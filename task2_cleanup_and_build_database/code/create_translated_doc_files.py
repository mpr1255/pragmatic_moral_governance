#%%
import sqlite3
import asyncio
import os
from pathlib import Path
from typing import List, Dict
from pydantic import create_model, Field
from openai import AsyncOpenAI
from docx import Document
from datetime import datetime
import logging
import json
from tenacity import retry, stop_after_attempt, wait_exponential
import time

# OpenAI client setup
client = AsyncOpenAI()

# Database and output configuration
TASK2_DIR = Path(__file__).resolve().parents[2]
REPO_ROOT = Path(__file__).resolve().parents[3]
DB_PATH = Path(os.getenv('WENMING_DB_PATH', TASK2_DIR / 'out' / 'concept_scores.db'))
OUTPUT_DIR = Path(os.getenv('WENMING_TRANSLATIONS_DIR', REPO_ROOT / 'task3_analyse_data' / 'out' / 'translations'))

# Ensure output directory exists
OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

# Semaphore for API calls
semaphore = asyncio.Semaphore(20)

def enable_wal_mode(conn):
    conn.execute('PRAGMA journal_mode=WAL')
    conn.execute('PRAGMA synchronous=NORMAL')
    conn.execute('PRAGMA cache_size=-64000')  # 64MB cache
    conn.execute('PRAGMA temp_store=MEMORY')

def get_db_connection():
    conn = sqlite3.connect(str(DB_PATH), check_same_thread=False)
    enable_wal_mode(conn)
    conn.row_factory = sqlite3.Row
    return conn

#%%
def get_random_rows(num_rows: int = 1) -> List[Dict]:
    with get_db_connection() as conn:
        cursor = conn.cursor()
        cursor.execute(f"SELECT * FROM file_index ORDER BY RANDOM() LIMIT {num_rows}")
        return [dict(row) for row in cursor.fetchall()]

def create_translation_model(start: int, end: int):
    fields = {str(i): (str | None, Field(default=None, description=f"Translated text for line {i}")) for i in range(start, end + 1)}
    return create_model('DynamicTranslationResponse', **fields)

@retry(stop=stop_after_attempt(3), wait=wait_exponential(multiplier=1, min=4, max=10))
async def translate_chunk(chunk: List[str], start: int, end: int) -> Dict[str, str]:
    DynamicTranslationResponse = create_translation_model(start, end)
    
    logging.debug(f"Translating chunk: start={start}, end={end}")
    logging.debug(f"Input chunk: {json.dumps(chunk, ensure_ascii=False)}")

    async with semaphore:
        messages = [
            {"role": "system", "content": f"You are a professional, accurate, and careful Chinese to English translator. You return ONLY English. Translate the following text from Chinese to English. Include all line numbers from {start} to {end}, even if the line is empty or just punctuation. Use JSON."},
            {"role": "user", "content": "\n".join(f"{i+start}. {line}" for i, line in enumerate(chunk))},
        ]
        logging.debug(f"Messages sent to API: {json.dumps(messages, ensure_ascii=False)}")

        completion = await client.chat.completions.create(
            model="gpt-4o-mini",
            messages=messages,
            response_format={"type": "json_object"},
        )

        logging.debug(f"Raw API response: {json.dumps(completion.model_dump(), ensure_ascii=False)}")

    try:
        response_content = completion.choices[0].message.content
        logging.debug(f"Response content: {response_content}")

        response = DynamicTranslationResponse.model_validate_json(response_content)
        logging.debug(f"Validated response: {json.dumps(response.model_dump(), ensure_ascii=False)}")

        return response.model_dump()
    except Exception as e:
        logging.error(f"Error processing response: {str(e)}")
        logging.error(f"Raw response content: {completion.choices[0].message.content}")
        raise  # This will trigger a retry

def create_translation_column():
    with get_db_connection() as conn:
        cursor = conn.cursor()
        cursor.execute("PRAGMA table_info(file_index)")
        columns = [col[1] for col in cursor.fetchall()]
        if 'translated_content' not in columns:
            cursor.execute("ALTER TABLE file_index ADD COLUMN translated_content TEXT")
        conn.commit()

def save_translation_to_db(row_id: str, translated_content: str):
    with get_db_connection() as conn:
        cursor = conn.cursor()
        cursor.execute("UPDATE file_index SET translated_content = ? WHERE id = ?", (translated_content, row_id))
        conn.commit()

@retry(stop=stop_after_attempt(3), wait=wait_exponential(multiplier=1, min=4, max=10))
async def process_document(content: str, chunk_size: int = 20) -> Dict[str, str]:
    lines = content.split('\n')
    all_translations = {}
    
    tasks = []
    for i in range(0, len(lines), chunk_size):
        chunk = lines[i:i+chunk_size]
        start = i + 1
        end = i + len(chunk)
        tasks.append(translate_chunk(chunk, start, end))
    
    chunks = await asyncio.gather(*tasks)
    for chunk in chunks:
        all_translations.update(chunk)
    
    # Ensure all keys are present and have string values
    for i in range(1, len(lines) + 1):
        if str(i) not in all_translations or all_translations[str(i)] is None:
            all_translations[str(i)] = f"[Translation missing for line {i}]"
    
    return all_translations

def create_word_document(row: Dict, translations: Dict[str, str], filename: Path):
    doc = Document()
    
    # Add metadata table
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    for field, value in row.items():
        if field != 'content':
            cells = table.add_row().cells
            cells[0].text = field.capitalize()
            cells[1].text = str(value)

    doc.add_paragraph()

    # Add content with translations
    for line_num, (original, translation) in enumerate(zip(row['content'].split('\n'), translations.values()), 1):
        p = doc.add_paragraph()
        p.add_run(f"{line_num}. ").bold = True
        p.add_run(f"{original}\n")
        p.add_run(f"   {translation}\n")

    doc.save(str(filename))

async def main(num_documents: int = 45):
    create_translation_column()
    random_rows = get_random_rows(num_documents)
    
    for i, row in enumerate(random_rows, 1):
        print(f"Processing document {i}/{num_documents}: {row['title']}")
        
        # Extract date from publish field
        try:
            publish_date = datetime.strptime(row['publish'], '%Y-%m-%d %H:%M:%S').strftime('%Y-%m-%d')
        except ValueError:
            publish_date = row['publish'] if row['publish'] else 'undated'
        
        # Create filename
        filename = f"{row['title']}_{row['office']}_{publish_date}.docx"
        filename = "".join(c for c in filename if c.isalnum() or c in ['_', '.', '-'])  # Remove any invalid characters
        full_path = OUTPUT_DIR / filename
        
        if full_path.exists():
            print(f"File {full_path} already exists. Skipping translation.")
            continue
        
        max_retries = 3
        for attempt in range(max_retries):
            try:
                translations = await process_document(row['content'])
                
                # Combine translations into a single string, preserving line breaks
                translated_content = '\n'.join(translations[str(i)] for i in range(1, len(translations) + 1))
                
                # Save translation to database
                save_translation_to_db(row['id'], translated_content)
                
                create_word_document(row, translations, full_path)
                print(f"Created {full_path}")
                print(f"Saved translation to database for document ID: {row['id']}")
                break  # Success, exit the retry loop
            except Exception as e:
                if attempt < max_retries - 1:
                    print(f"Error processing document {row['title']}, attempt {attempt + 1}/{max_retries}. Retrying...")
                    time.sleep(5)  # Wait 5 seconds before retrying
                else:
                    print(f"Failed to process document {row['title']} after {max_retries} attempts. Error: {str(e)}")
                    logging.error(f"Failed to process document {row['title']}. Error: {str(e)}")

#%%
if __name__ == "__main__":
    # Set up logging
    logging.basicConfig(filename='/tmp/wenming_translations.log', level=logging.DEBUG, 
                        format='%(asctime)s - %(levelname)s - %(message)s')

    # You can change the number of documents to process here
    asyncio.run(main(num_documents=45))

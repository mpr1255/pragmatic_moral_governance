import sqlite3
import os
import json
import argparse
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.style import WD_STYLE_TYPE
from bs4 import BeautifulSoup
from datetime import datetime
import markdown
from typing import Dict  # Add this import

# Default locations relative to repository structure
TASK2_DIR = Path(__file__).resolve().parents[2]
REPO_ROOT = Path(__file__).resolve().parents[3]
DEFAULT_DB_PATH = Path(os.getenv('WENMING_DB_PATH', TASK2_DIR / 'out' / 'concept_scores.db'))
DEFAULT_OUTPUT_DIR = Path(os.getenv('WENMING_EXPORT_DIR', REPO_ROOT / 'task3_analyse_data' / 'out' / 'exported_docs'))

# Ensure output directory exists
DEFAULT_OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

def get_db_connection(db_path):
    conn = sqlite3.connect(str(db_path))
    conn.row_factory = sqlite3.Row
    return conn

def get_rows(conn, query):
    cursor = conn.cursor()
    cursor.execute(query)
    return [dict(row) for row in cursor.fetchall()]

def create_markdown_document(row, columns_to_export):
    md_content = ""
    for column in columns_to_export:
        if column != 'content':
            value = row.get(column, '')
            if isinstance(value, str) and value.startswith('{') and value.endswith('}'):
                try:
                    json_value = json.loads(value)
                    md_content += f"## {column.capitalize()}\n\n"
                    md_content += json.dumps(json_value, indent=2, ensure_ascii=False)
                    md_content += "\n\n"
                except json.JSONDecodeError:
                    md_content += f"**{column.capitalize()}**: {value}\n\n"
            else:
                md_content += f"**{column.capitalize()}**: {value}\n\n"

    if 'content' in columns_to_export:
        md_content += "## Content\n\n"
        md_content += row.get('content', '')

    return md_content

def create_word_document(row: Dict, output_path: str):
    doc = Document()
    
    # Add metadata table
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    for field, value in row.items():
        if field not in ['content', 'all_five_concepts']:
            cells = table.add_row().cells
            cells[0].text = field.capitalize()
            cells[1].text = str(value)

    doc.add_paragraph()

    # Handle all_five_concepts
    if 'all_five_concepts' in row:
        doc.add_heading('All Five Concepts', level=2)
        try:
            data = json.loads(row['all_five_concepts'])
            table = doc.add_table(rows=1, cols=2)
            table.style = 'Table Grid'
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Key'
            hdr_cells[1].text = 'Value'
            for key, value in data.items():
                row_cells = table.add_row().cells
                row_cells[0].text = key
                row_cells[1].text = str(value)
            table.autofit = True
        except json.JSONDecodeError:
            doc.add_paragraph(row['all_five_concepts'])

    doc.add_paragraph()

    # Add content
    if 'content' in row:
        doc.add_heading('Content', level=2)
        for line in row['content'].split('\n'):
            doc.add_paragraph(line)

    doc.save(output_path)

def main(db_path, output_dir, table, out_cols, essential_col, query):
    output_dir = Path(output_dir)
    output_dir.mkdir(parents=True, exist_ok=True)
    
    with get_db_connection(db_path) as conn:
        rows = get_rows(conn, query)
    
    for i, row in enumerate(rows, 1):
        if row.get(essential_col):
            print(f"Processing document {i}/{len(rows)}: {row['title']}")
            
            publish_date = row.get('publish', 'undated')
            if publish_date != 'undated':
                try:
                    publish_date = datetime.strptime(publish_date, '%Y-%m-%d %H:%M:%S').strftime('%Y-%m-%d')
                except ValueError:
                    pass
            
            filename = f"{row['title']}_{row.get('office', 'unknown')}_{publish_date}.docx"
            filename = "".join(c for c in filename if c.isalnum() or c in ['_', '.', '-'])
            full_path = output_dir / filename
            
            if full_path.exists():
                print(f"Warning: Overwriting existing file {full_path}")
                
            markdown_content = create_markdown_document(row, out_cols)
            create_word_document(row, str(full_path))
            print(f"Created/Updated {full_path}")
        else:
            print(f"Skipping document {i}/{len(rows)}: {row['title']} (missing {essential_col})")

if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="Export data from SQLite to Word documents")
    parser.add_argument("--db", default=str(DEFAULT_DB_PATH), required=False, help="Path to the SQLite database")
    parser.add_argument("--output-dir", default=str(DEFAULT_OUTPUT_DIR), required=False, help="Output directory for Word documents")
    parser.add_argument("--table", required=False, help="Table name to query")
    parser.add_argument("--out-cols", required=True, help="Comma-separated list of columns to export")
    parser.add_argument("--essential-col", required=True, help="Column that must be non-null for export")
    parser.add_argument("--query", required=True, help="SQL query to select rows")

    args = parser.parse_args()

    out_cols = args.out_cols.split(',')
    main(args.db, args.output_dir, args.table, out_cols, args.essential_col, args.query)
